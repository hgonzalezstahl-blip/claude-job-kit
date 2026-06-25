#!/usr/bin/env python
"""
Quality-check tailored resume / cover .docx files before delivery.

Usage:
  py tools/verify_docx.py "<path to a .docx OR a folder>"

Checks each Resume*.docx / Cover*.docx for:
  - OVER-BOLD body bullets  (a whole bullet set to bold -> reads as AI clutter)
  - em-dashes (the clearest AI tell; use en-dash U+2013 or commas instead)
  - leftover placeholders ([COMPANY], {{ }}, XXXX, lorem, TODO)

Exit code 0 = all clean. Non-zero = something to fix.
Generic: no personal data. Run fix_bold.py to auto-correct over-bold bullets.
"""
import sys, os, glob
try:
    from docx import Document
except ImportError:
    print("python-docx is not installed. Run:  py -m pip install python-docx")
    sys.exit(2)

PLACEHOLDERS = ["{{", "}}", "[company]", "[role]", "[name]", "xxxx", "lorem", "todo", "placeholder", "<insert"]

def is_allcaps(t):
    letters = [c for c in t if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)

def fully_bold(p):
    runs = [r for r in p.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)

def check(path):
    name = os.path.basename(path)
    is_cover = name.lower().startswith("cover")
    try:
        d = Document(path)
    except Exception as e:
        return [f"CANNOT OPEN ({e})"]
    full = "\n".join(p.text for p in d.paragraphs)
    low = full.lower()
    problems = []
    # over-bold body bullets: fully-bold, long (>70 chars), after the first ALL-CAPS section header
    seen_header = False
    overbold = 0
    for p in d.paragraphs:
        t = p.text.strip()
        if is_allcaps(t):
            seen_header = True
            continue
        if seen_header and t and len(t) > 70 and fully_bold(p):
            overbold += 1
    if overbold:
        problems.append(f"{overbold} OVER-BOLD bullet(s) -> run fix_bold.py")
    # covers should have no fully-bold body paragraph except a short 'Re:' line
    if is_cover:
        bb = 0
        for p in d.paragraphs:
            t = p.text.strip()
            if t and len(t) > 40 and not t.lower().startswith(("re:", "hiring team")) and fully_bold(p):
                bb += 1
        if bb:
            problems.append(f"{bb} bold body paragraph(s) in cover (covers are plain prose)")
    if "—" in full:
        problems.append(f"{full.count(chr(8212))} em-dash(es) -> use en-dash or commas")
    hits = [w for w in PLACEHOLDERS if w in low]
    if hits:
        problems.append(f"placeholder(s) left: {hits}")
    return problems

def main():
    if len(sys.argv) < 2:
        print(__doc__); sys.exit(2)
    target = sys.argv[1]
    if os.path.isdir(target):
        files = glob.glob(os.path.join(target, "**", "Resume*.docx"), recursive=True) + \
                glob.glob(os.path.join(target, "**", "Cover*.docx"), recursive=True)
    else:
        files = [target]
    if not files:
        print("No Resume*.docx / Cover*.docx found at:", target); sys.exit(2)
    any_bad = False
    for f in sorted(files):
        probs = check(f)
        tag = "OK  " if not probs else "FAIL"
        if probs: any_bad = True
        print(f"[{tag}] {os.path.basename(f)}" + ("" if not probs else "  -> " + "; ".join(probs)))
    print("\nAll clean." if not any_bad else "\nSome files need fixing (see above).")
    sys.exit(0 if not any_bad else 1)

if __name__ == "__main__":
    main()
